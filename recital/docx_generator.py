from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ai.models import StudentPerformance


def generate_recital_docx(
    performances: List[StudentPerformance],
    studio_name: str,
    recital_title: str,
    recital_date: str,
    accompanist: str,
    footer_text: str,
    output_dir: str = "output",
    filename: str = "recital_program.docx",
):
    """
    Generate a strictly formatted Suzuki recital program DOCX.

    CRITICAL GUARANTEES:
    - Every student/piece generates a table, even if unmatched
    - Unmatched fields are rendered as blank (never omitted)
    - Formatting follows the original prompt exactly

    Args:
        performances: List of student performances
        studio_name: Name of the studio (H2)
        recital_title: Title of the recital (H1)
        recital_date: Date/time of the recital
        accompanist: Name of the accompanist
        footer_text: Footer text for acknowledgments
        output_dir: Directory to save the file
        filename: Name of the output file
    """

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ============================================================
    # HEADER (exactly as specified)
    # ============================================================

    # Studio Name (bold, 14pt, centered)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(studio_name)
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(6)

    # Recital Title (bold, 18pt, centered)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(recital_title)
    run.bold = True
    run.font.size = Pt(18)
    p.paragraph_format.space_after = Pt(6)

    # Date / time (centered)
    p = doc.add_paragraph(recital_date)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)

    # Body line (centered, paragraph style)
    p = doc.add_paragraph(
        f"Students of Dr. Ryan Graebert\nAccompanied by {accompanist}"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    # ============================================================
    # PROGRAM ENTRIES
    # ============================================================

    for student in performances:
        # Sort student's pieces by library_index
        sorted_pieces = sorted(
            student.pieces,
            key=lambda p: (
                (999, 9999)
                if not p.library_index
                else tuple(map(int, p.library_index.split(".")))
            ),
        )

        num_pieces = len(sorted_pieces)

        # ---- 2-column table, no borders, 3/5 and 2/5 width ----
        # Create table with: num_pieces rows for pieces + 1 row for student name
        table = doc.add_table(rows=num_pieces + 1, cols=2)
        table.autofit = False
        table.allow_autofit = False

        # Set column widths: 3/5 for title, 2/5 for composer (within margins)
        # Using 6 inches total width to stay within margins
        table.columns[0].width = Inches(3.6)  # 3/5 of 6 inches
        table.columns[1].width = Inches(2.4)  # 2/5 of 6 inches

        # Add each piece as a row
        for piece_idx, piece in enumerate(sorted_pieces):
            # ====================================================
            # Left Cell: Title + Movements
            # ====================================================
            left_cell = table.cell(piece_idx, 0)
            left_p = left_cell.paragraphs[0]
            left_p.paragraph_format.space_after = Pt(0)  # Remove space after title

            # Title (BOLD if present, blank if not)
            title_text = piece.title or ""
            title_run = left_p.add_run(title_text)
            title_run.bold = True

            # Movements (each on its own indented line, no extra spacing)
            if piece.movements:
                for movement in piece.movements:
                    mp = left_cell.add_paragraph(movement)
                    mp.paragraph_format.left_indent = Inches(0.3)
                    mp.paragraph_format.space_after = Pt(
                        0
                    )  # No whitespace between movements
            else:
                # IMPORTANT: preserve spacing even when no movements
                left_cell.add_paragraph("")

            # ====================================================
            # Right Cell: Composer (right-aligned)
            # ====================================================
            right_cell = table.cell(piece_idx, 1)
            right_p = right_cell.paragraphs[0]
            right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_p.paragraph_format.space_after = Pt(0)  # Remove space after composer

            composer_text = piece.composer or ""
            right_p.add_run(composer_text)

        # ====================================================
        # Last Row: Student Name (merged, centered)
        # ====================================================
        bottom_cell = table.cell(num_pieces, 0)
        bottom_cell.merge(table.cell(num_pieces, 1))

        sp = bottom_cell.paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.add_run(student.student_name)

        # Spacer paragraph between students
        doc.add_paragraph("")

    # ============================================================
    # FOOTER (exactly as specified)
    # ============================================================

    p = doc.add_paragraph(footer_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============================================================
    # SAVE FILE
    # ============================================================

    full_path = output_path / filename
    doc.save(full_path)

    return full_path
